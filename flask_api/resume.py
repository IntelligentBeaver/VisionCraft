# -*- coding: utf-8 -*-
"""resume.ipynb

Automatically generated by Colab.

Original file is located at
    https://colab.research.google.com/drive/1-Olq0UZjqzUjohvYzGFn5h_SXlqK__Nm
"""

# !pip install pdfplumber pandas flask-ngrok

import pdfplumber
import os
from google.colab import files

# Upload a PDF
uploaded = files.upload()

# Get the filename
pdf_filename = list(uploaded.keys())[0]

# Extract text
def extract_text_from_pdf(pdf_path):
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    return text

resume_text = extract_text_from_pdf(pdf_filename)
print(resume_text[:1000])  # Print first 1000 characters for preview

import re

def extract_sections(text):
    sections = {
        "Name": "",
        "Contact Info": "",
        "Summary":"",
        "Work Experience": "",
        "Education": "",
        "Skills": "",
        "Certifications": "",
    }

    # Name Extraction (First line assumption)
    lines = text.split("\n")
    if len(lines) > 0:
        sections["Name"] = lines[0] if len(lines[0].split()) <= 4 else "[Name Not Detected]"

    # Contact Info (Detect Email, Phone)
    email_match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
    phone_match = re.search(r"\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}", text)

    if email_match:
        sections["Contact Info"] += f"Email: {email_match.group(0)}\n"
    else:
        sections["Contact Info"] += "Email: [Not Found]\n"

    if phone_match:
        sections["Contact Info"] += f"Phone: {phone_match.group(0)}\n"
    else:
        sections["Contact Info"] += "Phone: [Not Found]\n"

    # Section Detection (Regex-based)
    patterns = {
        "Summary": r"(?i)\b(summary|profile|overview)\b[\s\S]*?\n\n" ,
        "Work Experience": r"(?i)work experience|employment history|professional experience",
        "Education": r"(?i)education|academic background|degrees",
        "Skills": r"(?i)skills|technical skills|competencies",
        "Certifications": r"(?i)certifications|licenses|courses",
    }

    for key, pattern in patterns.items():
        match = re.search(pattern, text)
        if match:
            sections[key] = text[match.start(): match.start()+500]  # Extract a portion of text

    return sections

resume_sections = extract_sections(resume_text)
print(resume_sections)

def check_missing_sections(sections):
    missing_sections = [key for key, value in sections.items() if value == ""]

    if missing_sections:
        print("\n🚨 Missing Sections Found:")
        for sec in missing_sections:
            print(f"- {sec} is missing. Please add it manually.")
    else:
        print("\n✅ Resume is complete!")

check_missing_sections(resume_sections)

# !pip install python-docx
from docx import Document

def generate_ats_friendly_resume(sections, filename="ATS_Resume.docx"):
    doc = Document()

    doc.add_heading("ATS-Friendly Resume", level=1)

    # Add each section
    for section, content in sections.items():
        doc.add_heading(section, level=2)
        if content:
            doc.add_paragraph(content)
        else:
            doc.add_paragraph(f"[{section} Not Found - Please Add]")

    doc.save(filename)
    return filename

docx_filename = generate_ats_friendly_resume(resume_sections)
print(f"\n✅ ATS-Friendly Resume saved as {docx_filename}")