# from flask import Flask, request, jsonify, send_from_directory
# from flask_cors import CORS
# import pdfplumber
# import re
# import os
# import time
# from docx import Document
# import pickle
# from sklearn.metrics.pairwise import cosine_similarity
# import numpy as np

# # Load saved models
# with open("tfidf_vectorizer0.pkl", "rb") as file:
#     tfidf = pickle.load(file)

# with open("vector_matrix0.pkl", "rb") as file:
#     vector_matrix = pickle.load(file)

# with open("job_data.pkl", "rb") as file:
#     jobs_df = pickle.load(file)

# app = Flask(__name__)
# CORS(app)  # Enable CORS for Laravel API to communicate



# @app.route('/recommend', methods=['POST'])
# def recommend_jobs():
#     data = request.json
#     user_skills = data.get("skills", "")
#     user_industry = data.get("industry", "")
#     user_functional_area = data.get("functional_area", "")

#     if not user_skills or not user_industry or not user_functional_area:
#         return jsonify({"error": "Missing input fields"}), 400

#     # Combine user input
#     user_input = f"{user_skills} {user_industry} {user_functional_area}"

#     # Convert user input into vector
#     user_vector = tfidf.transform([user_input])

#     # Compute similarity
#     user_similarity_scores = cosine_similarity(user_vector, vector_matrix)

#     # Get top 5 recommendations
#     sorted_indices = user_similarity_scores[0].argsort()[::-1][:5]

#     recommendations = []
#     for idx in sorted_indices:
#         recommendations.append({
#             "Job Title": jobs_df.iloc[idx]['Job Title'],
#             "Industry": jobs_df.iloc[idx]['Industry'],
#             "Functional Area": jobs_df.iloc[idx]['Functional Area'],
#             "Role": jobs_df.iloc[idx]['Role'],  
#             "Similarity Score": float(user_similarity_scores[0][idx])
#         })

#     return jsonify({"recommendations": recommendations})
# # @app.route('/test', methods=['GET'])
# # def test_api():
# #     return jsonify({"message": "Hello, World!"})

# if __name__ == '__main__':
#     app.run(host="0.0.0.0", port=5000, debug=True)

from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
import pdfplumber
import re
import os
import time
from docx import Document
import pickle
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np

# -------------------------------
# Load Job Recommender AI Model
# -------------------------------
with open("tfidf_vectorizer0.pkl", "rb") as file:
    tfidf = pickle.load(file)

with open("vector_matrix0.pkl", "rb") as file:
    vector_matrix = pickle.load(file)

with open("job_data.pkl", "rb") as file:
    jobs_df = pickle.load(file)

# -------------------------------
# Initialize Flask App
# -------------------------------
app = Flask(__name__)
CORS(app)  # Enable CORS for Laravel API to communicate

UPLOAD_FOLDER = "saved_resumes"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)  # Ensure folder exists

# -------------------------------
# Resume Optimizer API
# -------------------------------

# Function to extract text from PDF
def extract_text_from_pdf(pdf_path):
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text.strip()

# Function to extract sections from resume
def extract_sections(text):
    sections = {
        "Name": "",
        "Contact Info": "",
        "Summary": "",
        "Work Experience": "",
        "Education": "",
        "Skills": "",
        "Certifications": "",
    }

    lines = text.split("\n")
    if len(lines) > 0:
        sections["Name"] = lines[0] if len(lines[0].split()) <= 4 else "[Name Not Detected]"

    email_match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
    phone_match = re.search(r"\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}", text)

    sections["Contact Info"] += f"Email: {email_match.group(0)}\n" if email_match else "Email: [Not Found]\n"
    sections["Contact Info"] += f"Phone: {phone_match.group(0)}\n" if phone_match else "Phone: [Not Found]\n"

    patterns = {
        "Summary": r"(?i)\b(summary|profile|overview)\b[\s\S]*?\n\n",
        "Work Experience": r"(?i)work experience|employment history|professional experience",
        "Education": r"(?i)education|academic background|degrees",
        "Skills": r"(?i)skills|technical skills|competencies",
        "Certifications": r"(?i)certifications|licenses|courses",
    }

    for key, pattern in patterns.items():
        match = re.search(pattern, text)
        if match:
            sections[key] = text[match.start(): match.start()+500]

    return sections

# Function to generate an ATS-friendly .docx
def generate_ats_friendly_resume(sections, filename):
    doc = Document()
    doc.add_heading("ATS-Friendly Resume", level=1)

    for section, content in sections.items():
        doc.add_heading(section, level=2)
        doc.add_paragraph(content if content else f"[{section} Not Found - Please Add]")

    doc.save(filename)
    return filename

@app.route('/upload_resume', methods=['POST'])
def upload_resume():
    file = request.files['resume']
    file_path = os.path.join(UPLOAD_FOLDER, "uploaded_resume.pdf")
    file.save(file_path)

    text = extract_text_from_pdf(file_path)
    sections = extract_sections(text)

    ats_filename = f"ATS_Resume_{int(time.time())}.docx"
    ats_file_path = os.path.join(UPLOAD_FOLDER, ats_filename)
    generate_ats_friendly_resume(sections, ats_file_path)

    return jsonify({
        "message": "Resume processed successfully!",
        "filename": ats_filename,
        "download_url": f"/download/{ats_filename}"
    })

@app.route('/download/<filename>', methods=['GET'])
def download_resume(filename):
    return send_from_directory(UPLOAD_FOLDER, filename, as_attachment=True)

# -------------------------------
# Job Recommender API
# -------------------------------

@app.route('/recommend', methods=['POST'])
def recommend_jobs():
    data = request.json
    user_skills = data.get("skills", "")
    user_industry = data.get("industry", "")
    user_functional_area = data.get("functional_area", "")

    if not user_skills or not user_industry or not user_functional_area:
        return jsonify({"error": "Missing input fields"}), 400

    user_input = f"{user_skills} {user_industry} {user_functional_area}"

    user_vector = tfidf.transform([user_input])

    user_similarity_scores = cosine_similarity(user_vector, vector_matrix)

    sorted_indices = user_similarity_scores[0].argsort()[::-1][:5]

    recommendations = []
    for idx in sorted_indices:
        recommendations.append({
            "Job Title": jobs_df.iloc[idx]['Job Title'],
            "Industry": jobs_df.iloc[idx]['Industry'],
            "Functional Area": jobs_df.iloc[idx]['Functional Area'],
            "Role": jobs_df.iloc[idx]['Role'],  
            "Similarity Score": float(user_similarity_scores[0][idx])
        })

    return jsonify({"recommendations": recommendations})

if __name__ == '__main__':
    app.run(host="0.0.0.0", port=5000, debug=True)